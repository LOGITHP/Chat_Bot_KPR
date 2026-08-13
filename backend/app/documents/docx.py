import docx
from typing import List
from app.documents.base import BaseParser, DocumentChunk

class DOCXParser(BaseParser):
    def parse(self, file_path: str) -> List[DocumentChunk]:
        chunks = []
        try:
            doc = docx.Document(file_path)
            current_section = None
            
            for para in doc.paragraphs:
                if not para.text.strip():
                    continue
                    
                if para.style.name.startswith('Heading'):
                    current_section = para.text.strip()
                
                chunks.append(DocumentChunk(
                    text=para.text.strip(),
                    section=current_section
                ))
            
            for table in doc.tables:
                for i, row in enumerate(table.rows):
                    row_data = [cell.text.strip() for cell in row.cells]
                    text = " | ".join(row_data)
                    if text.strip():
                        chunks.append(DocumentChunk(
                            text=text,
                            section="Table",
                            row=i + 1
                        ))
            return chunks
        except Exception as e:
            print(f"Error parsing DOCX: {e}")
            return []
